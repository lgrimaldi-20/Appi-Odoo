"""
Genera el informe de estructura, flujo y estado de la integracion Smartier-Odoo.

Documento .docx para la directiva y el equipo tecnico: arquitectura del
middleware, esquema grafico del flujo, resultado de las pruebas y el inventario
de campos que faltan para poder facturar.

El esquema del flujo se dibuja con TABLAS de Word (no imagenes), asi que se ve
igual en Word, Google Docs y LibreOffice, y puede editarse en el propio
documento.

Requiere python-docx, que NO esta en requirements.txt a proposito: la API no lo
necesita para funcionar, solo este script.
    pip install python-docx

Uso:  python scripts/generar_informe_estructura.py
El .docx resultante esta en .gitignore: es un entregable, se regenera con esto.
"""

import re

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# Paleta: azul tecnico tomado del propio documento del cliente.
AZUL = RGBColor(0x1F, 0x4E, 0x79)
TINTA = RGBColor(0x16, 0x20, 0x2B)
TENUE = RGBColor(0x5A, 0x66, 0x72)
GRANATE = RGBColor(0x9B, 0x2C, 0x2C)
AMBAR = RGBColor(0x8A, 0x50, 0x0A)
VERDE = RGBColor(0x2F, 0x6B, 0x4F)

# Fondos (hex sin almohadilla, como los quiere OOXML).
F_AZUL = "EAF1F8"
F_GRANATE = "FAEDED"
F_AMBAR = "FBF1E0"
F_VERDE = "E8F2EC"
F_GRIS = "F2F1EE"

doc = Document()

for s in doc.sections:
    s.top_margin = Cm(2.0)
    s.bottom_margin = Cm(2.0)
    s.left_margin = Cm(2.2)
    s.right_margin = Cm(2.2)

normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
normal.font.color.rgb = TINTA
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.15


# --- utilidades de formato -------------------------------------------------

def sombrear(parrafo, color_hex):
    """Fondo de color en un parrafo (python-docx no lo expone)."""
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), color_hex)
    parrafo._p.get_or_add_pPr().append(shd)


def sombrear_celda(celda, color_hex):
    """Fondo de color en una celda de tabla."""
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), color_hex)
    celda._tc.get_or_add_tcPr().append(shd)


def borde(parrafo, lado, color_hex, ancho=18):
    pbdr = parrafo._p.get_or_add_pPr().find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        parrafo._p.get_or_add_pPr().append(pbdr)
    el = OxmlElement(f"w:{lado}")
    el.set(qn("w:val"), "single")
    el.set(qn("w:sz"), str(ancho))
    el.set(qn("w:space"), "6")
    el.set(qn("w:color"), color_hex)
    pbdr.append(el)


def sin_bordes(tabla):
    """Quita todos los bordes de una tabla (para los esquemas)."""
    tbl_pr = tabla._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for lado in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{lado}")
        el.set(qn("w:val"), "none")
        borders.append(el)
    tbl_pr.append(borders)


def texto(p, contenido, negrita=False, mono=False, color=None,
          cursiva=False, tam=None):
    run = p.add_run(contenido)
    run.bold = negrita
    run.italic = cursiva
    if mono:
        run.font.name = "Consolas"
        run.font.size = Pt(9.5)
    if tam:
        run.font.size = Pt(tam)
    if color:
        run.font.color.rgb = color
    return run


def rico(p, contenido):
    """Escribe interpretando **negrita** y `monoespaciada`, incluso anidadas."""
    def escribir(fragmento, negrita=False):
        for sub in re.split(r"(`[^`]+`)", fragmento):
            if not sub:
                continue
            if sub.startswith("`") and sub.endswith("`"):
                texto(p, sub[1:-1], mono=True, color=AZUL, negrita=negrita)
            else:
                texto(p, sub, negrita=negrita)

    for trozo in re.split(r"(\*\*.+?\*\*)", contenido):
        if not trozo:
            continue
        if trozo.startswith("**") and trozo.endswith("**"):
            escribir(trozo[2:-2], negrita=True)
        else:
            escribir(trozo)


def parrafo(contenido, indent=0.0, espacio=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(espacio)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    rico(p, contenido)
    return p


def seccion(numero, titulo, nota=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(3)
    texto(p, f"{numero}  ", negrita=True, mono=True, color=AZUL, tam=11)
    run = texto(p, titulo, negrita=True, tam=15)
    run.font.name = "Cambria"
    borde(p, "bottom", "16202B", 12)
    if nota:
        q = doc.add_paragraph()
        q.paragraph_format.space_after = Pt(8)
        texto(q, nota, cursiva=True, color=TENUE, tam=10)


def subtitulo(titulo):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(3)
    texto(p, titulo, negrita=True, tam=11.5, color=AZUL)
    return p


def aviso(titulo, frases, tono="granate"):
    fondo, color, hexb = {
        "granate": (F_GRANATE, GRANATE, "9B2C2C"),
        "ambar": (F_AMBAR, AMBAR, "B4690E"),
        "azul": (F_AZUL, AZUL, "1F4E79"),
        "verde": (F_VERDE, VERDE, "2F6B4F"),
    }[tono]
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.3)
    sombrear(p, fondo)
    borde(p, "left", hexb)
    texto(p, titulo.upper(), negrita=True, color=color, tam=8.5)
    for f in frases:
        q = doc.add_paragraph()
        q.paragraph_format.left_indent = Cm(0.3)
        q.paragraph_format.space_after = Pt(4)
        sombrear(q, fondo)
        borde(q, "left", hexb)
        rico(q, f)


def viñetas(items, indent=0.75):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Cm(indent)
        p.paragraph_format.space_after = Pt(3)
        rico(p, it)


def tabla_datos(encabezados, filas, anchos=None, colores_col=None):
    t = doc.add_table(rows=1, cols=len(encabezados))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, cab in enumerate(encabezados):
        celda = t.rows[0].cells[i]
        celda.text = ""
        p = celda.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        texto(p, cab.upper(), negrita=True, color=AZUL, tam=8)
        sombrear_celda(celda, F_AZUL)
    for fila in filas:
        celdas = t.add_row().cells
        for i, valor in enumerate(fila):
            celdas[i].text = ""
            p = celdas[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            col = (colores_col or {}).get(i)
            if callable(col):
                color, negrita = col(valor)
                texto(p, valor, color=color, negrita=negrita, tam=9)
            else:
                rico(p, valor)
                for r in p.runs:
                    r.font.size = Pt(9)
    if anchos:
        for fila in t.rows:
            for i, ancho in enumerate(anchos):
                fila.cells[i].width = Cm(ancho)
    return t


def caja_esquema(celda, titulo, lineas, fondo, color_titulo, hex_borde):
    """Dibuja una 'caja' del esquema dentro de una celda de tabla."""
    celda.text = ""
    sombrear_celda(celda, fondo)
    p = celda.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(1)
    texto(p, titulo, negrita=True, color=color_titulo, tam=9)
    for linea, mono in lineas:
        q = celda.add_paragraph()
        q.alignment = WD_ALIGN_PARAGRAPH.CENTER
        q.paragraph_format.space_after = Pt(1)
        texto(q, linea, mono=mono, tam=7.5,
              color=TENUE if not mono else AZUL)
    celda.paragraphs[-1].paragraph_format.space_after = Pt(4)


def flecha(celda, etiqueta=None, simbolo="→"):
    celda.text = ""
    p = celda.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(0)
    texto(p, simbolo, negrita=True, color=AZUL, tam=16)
    if etiqueta:
        q = celda.add_paragraph()
        q.alignment = WD_ALIGN_PARAGRAPH.CENTER
        q.paragraph_format.space_after = Pt(0)
        texto(q, etiqueta, tam=7, color=TENUE)


def pie_figura(texto_pie):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    texto(p, texto_pie, cursiva=True, color=TENUE, tam=8.5)


# ===========================================================================
# PORTADA
# ===========================================================================
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(2)
texto(p, "TURICOPY · INTEGRACIÓN SMARTIER–ODOO", negrita=True, color=AZUL, tam=9)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(4)
run = texto(p, "Estructura, flujo y estado de la integración", negrita=True, tam=23)
run.font.name = "Cambria"

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(10)
texto(p, "Informe técnico del middleware construido: cómo está organizado, por "
         "dónde circulan los datos, qué se ha probado contra los sistemas reales "
         "y qué información falta para poder emitir la primera factura.",
      color=TENUE, tam=11)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(12)
texto(p, "Smart Automata", negrita=True, tam=9)
texto(p, "   ·   24 de agosto de 2026", tam=9, color=TENUE)
texto(p, "   ·   Rama ", tam=9, color=TENUE)
texto(p, "Turicopy-V17", mono=True, tam=9)
texto(p, "   ·   109 pruebas automáticas en verde", tam=9, color=TENUE)

aviso("Situación en una frase", [
    "El middleware está **construido y probado**; lo que impide facturar no es "
    "el software, sino que **faltan datos en el origen**: ningún cliente de "
    "Smartier tiene RIF, y todavía no existe ninguna nota de entrega sobre la "
    "que facturar.",
], tono="ambar")


# ===========================================================================
# 01 · RESUMEN EJECUTIVO
# ===========================================================================
seccion("01", "Resumen ejecutivo")

parrafo(
    "Turicopy usa **Smartier** para pedidos y producción, y **Odoo** para la "
    "contabilidad. El middleware que hemos construido conecta ambos: lee de la "
    "API de Smartier, traduce los datos al formato contable y crea la factura "
    "en Odoo, sin que nadie tenga que copiar información a mano."
)

subtitulo("Qué está terminado")
viñetas([
    "**Un solo servicio** cubre todo el recorrido: lectura de Smartier, cola "
    "intermedia y sincronización con Odoo. No hacen falta dos middlewares.",
    "**Sin duplicados por diseño.** Reenviar la misma nota no crea una segunda "
    "factura: está verificado contra Odoo real.",
    "**Panel de control** para ver qué entró, qué se facturó y qué falló, sin "
    "leer registros técnicos.",
    "**Facturación en bolívares y en dólares**, con la conversión resuelta por "
    "Odoo a la tasa vigente.",
])

subtitulo("Qué falta, y no depende de nosotros")
viñetas([
    "**El RIF de los clientes.** Los 3 clientes de Smartier lo tienen vacío. "
    "Sin identificación fiscal, Odoo rechaza la factura.",
    "**Notas de entrega.** No existe ninguna todavía, así que el tramo final "
    "no se ha podido probar con datos reales.",
    "**Una instancia de Odoo activa.** La de pruebas era temporal y se eliminó.",
])


# ===========================================================================
# 02 · ESQUEMA DEL FLUJO
# ===========================================================================
seccion("02", "Esquema del flujo",
        "Por dónde circula un dato desde que se entrega la mercancía hasta que "
        "existe la factura.")

subtitulo("Vista general")

t = doc.add_table(rows=1, cols=5)
sin_bordes(t)
t.alignment = WD_TABLE_ALIGNMENT.CENTER
c = t.rows[0].cells

caja_esquema(c[0], "SMARTIER", [
    ("Pedidos y producción", False),
    ("API de solo lectura", False),
    ("5 peticiones/seg", True),
], F_GRIS, TINTA, "5A6672")
flecha(c[1], "lee cada 5 min")
caja_esquema(c[2], "MIDDLEWARE", [
    ("Traduce y controla", False),
    ("Un solo servicio", False),
    ("FastAPI + Celery", True),
], F_AZUL, AZUL, "1F4E79")
flecha(c[3], "crea factura")
caja_esquema(c[4], "ODOO", [
    ("Contabilidad y fiscal", False),
    ("Emite y contabiliza", False),
    ("JSON-RPC", True),
], F_VERDE, VERDE, "2F6B4F")

for fila in t.rows:
    for i, ancho in enumerate((4.2, 1.9, 4.2, 1.9, 4.2)):
        fila.cells[i].width = Cm(ancho)

pie_figura("Los datos van en un solo sentido: Smartier no recibe nada de vuelta, "
           "porque su API no permite escribir.")

subtitulo("Detalle interno: el middleware por dentro")

t = doc.add_table(rows=3, cols=5)
sin_bordes(t)
t.alignment = WD_TABLE_ALIGNMENT.CENTER

# Fila 1: ingesta
c = t.rows[0].cells
caja_esquema(c[0], "API SMARTIER", [
    ("notas-entrega", True),
    ("clientes · productos", True),
], F_GRIS, TINTA, "5A6672")
flecha(c[1], "1. lee lo nuevo")
caja_esquema(c[2], "INGESTA", [
    ("Traduce cada nota", False),
    ("al formato contable", False),
    ("ingesta_smartier.py", True),
], F_AZUL, AZUL, "1F4E79")
flecha(c[3], "2. encola")
caja_esquema(c[4], "COLA", [
    ("Base de datos", False),
    ("PENDIENTE", True),
    ("cola_sincronizacion", True),
], F_AMBAR, AMBAR, "B4690E")

# Fila 2: bajada visual hacia el poller
c = t.rows[1].cells
for celda in c:
    celda.text = ""
p = c[4].paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(2)
p.paragraph_format.space_after = Pt(2)
texto(p, "↓", negrita=True, color=AZUL, tam=14)
q = c[4].add_paragraph()
q.alignment = WD_ALIGN_PARAGRAPH.CENTER
texto(q, "3. toma un lote", tam=7, color=TENUE)

# Fila 3: poller -> Odoo
c = t.rows[2].cells
caja_esquema(c[0], "ODOO", [
    ("Factura emitida", False),
    ("y contabilizada", False),
], F_VERDE, VERDE, "2F6B4F")
flecha(c[1], "5. crea y valida", simbolo="←")
caja_esquema(c[2], "POLLER", [
    ("Comprueba duplicados", False),
    ("y sincroniza", False),
    ("poller.py", True),
], F_AZUL, AZUL, "1F4E79")
flecha(c[3], "4. procesa", simbolo="←")
caja_esquema(c[4], "REGISTRO", [
    ("PROCESADO / ERROR", True),
    ("visible en el panel", False),
], F_GRIS, TINTA, "5A6672")

for fila in t.rows:
    for i, ancho in enumerate((4.2, 1.9, 4.2, 1.9, 4.2)):
        fila.cells[i].width = Cm(ancho)

pie_figura("El recorrido atraviesa la cola: el mismo servicio escribe por un "
           "lado (ingesta) y lee por el otro (poller). El registro de lo hecho "
           "se guarda en el middleware, nunca en Smartier.")

aviso("Por qué existe la cola intermedia", [
    "**Absorbe los ritmos distintos.** Smartier limita a 5 peticiones por "
    "segundo; Odoo tiene su propia velocidad. La cola desacopla ambos.",
    "**Protege ante caídas.** Si Odoo no responde, lo ya leído espera en la "
    "cola en lugar de perderse.",
    "**Hace visible el problema.** Cuando una factura falla, se ve en qué "
    "punto exacto y por qué.",
], tono="azul")


# ===========================================================================
# 03 · ESTRUCTURA DEL SOFTWARE
# ===========================================================================
seccion("03", "Estructura del software",
        "Piezas que componen el middleware y responsabilidad de cada una.")

subtitulo("Módulos de la integración con Smartier")
tabla_datos(
    ("Módulo", "Responsabilidad"),
    [
        ("`smartier_client.py`",
         "Habla con la API: paginación, control del límite de 5 peticiones por "
         "segundo y reintentos si el servidor responde con error temporal."),
        ("`ingesta_smartier.py`",
         "Traduce cada nota de entrega al formato contable y la deja en la cola. "
         "Recuerda hasta dónde leyó para no repetir trabajo."),
        ("`poller.py`",
         "Toma lotes de la cola y los sincroniza con Odoo. Aísla los fallos: "
         "una nota con error no bloquea a las demás."),
        ("`sincronizador.py`",
         "Crea y contabiliza la factura en Odoo, comprobando antes que no exista "
         "ya. Es la garantía de que no se duplica."),
        ("`mappings.yaml`",
         "Correspondencia entre campos de origen y campos de Odoo. Se ajusta sin "
         "tocar el programa."),
        ("`observabilidad.py`",
         "Alimenta el panel de control con el estado de cada sincronización."),
    ],
    anchos=(4.6, 11.8),
)

subtitulo("Qué puede consultarse desde fuera")
tabla_datos(
    ("Dirección", "Para qué sirve"),
    [
        ("`/panel`", "Panel visual: estado de todo, en el navegador."),
        ("`/smartier/diagnostico`", "Radiografía de Smartier: cuántos registros hay y qué falta."),
        ("`/smartier/ingerir`", "Fuerza una lectura de Smartier sin esperar al turno automático."),
        ("`/poller/ejecutar`", "Fuerza el envío a Odoo de lo que haya en la cola."),
        ("`/health`", "Comprobación de que el servicio está vivo y conectado."),
    ],
    anchos=(5.0, 11.4),
)


# ===========================================================================
# 04 · PRUEBAS REALIZADAS
# ===========================================================================
seccion("04", "Pruebas realizadas",
        "Todo lo siguiente se ejecutó contra los sistemas reales, no en simulación.")

subtitulo("Contra la API de Smartier · 21 y 24 de agosto")
tabla_datos(
    ("Prueba", "Resultado"),
    [
        ("Conexión y autenticación", "Correcta con la clave facilitada"),
        ("Lectura de clientes", "3 registros"),
        ("Lectura de productos", "50 registros (eran 48 el día 21)"),
        ("Lectura de notas de entrega", "0 registros"),
        ("Lectura de órdenes", "0 registros"),
        ("Traducción al formato contable", "Correcta, comprobada con datos reales"),
        ("Diagnóstico automático", "Detecta por sí solo los clientes sin RIF"),
    ],
    anchos=(8.2, 8.2),
)

subtitulo("Contra Odoo 17 · antes de que se eliminara la instancia")
tabla_datos(
    ("Prueba", "Resultado"),
    [
        ("Crear factura con IVA", "Emitida y contabilizada"),
        ("Reenviar la misma factura", "No se duplicó: una sola en Odoo"),
        ("Factura en bolívares", "1.150,00 Bs, convertidos por Odoo a la tasa vigente"),
        ("Registrar un cobro", "Correcto"),
        ("Conciliar factura y cobro", "Saldo pendiente a cero"),
        ("Ajuste de existencias", "0 → 50 → 60 unidades"),
        ("Repetir un ajuste", "No se sumó dos veces: quedó en 60"),
    ],
    anchos=(8.2, 8.2),
)

aviso("Dos hallazgos de las pruebas", [
    "**La API acepta filtros que no existen.** Enviamos un parámetro inventado "
    "y respondió con normalidad en lugar de rechazarlo. Importa porque el "
    "control de duplicados previsto se apoya en un filtro (`TieneFactura`): si "
    "se escribiera mal, la sincronización traería también las notas ya "
    "facturadas, **sin ningún aviso**.",
    "**La documentación estaba incompleta.** El manual describía el cliente con "
    "solo dos campos; la API real devuelve seis, incluido el campo del RIF. "
    "Y los productos sí traen su porcentaje de IVA, algo que el manual no "
    "reflejaba.",
], tono="granate")


# ===========================================================================
# 05 · CAMPOS FALTANTES
# ===========================================================================
seccion("05", "Campos que faltan para poder facturar",
        "Comparación entre lo que Odoo exige y lo que Smartier entrega hoy.")


def color_estado(valor):
    if valor.startswith("Falta"):
        return GRANATE, True
    if valor.startswith("Lo da"):
        return VERDE, False
    return AMBAR, False


tabla_datos(
    ("Dato que Odoo necesita", "Situación", "Cómo se resuelve"),
    [
        ("RIF del cliente", "Falta", "Cargarlo en Smartier, o en Odoo una sola vez por cliente"),
        ("Notas de entrega", "Falta", "Que Smartier genere las primeras"),
        ("Número de control fiscal", "Falta", "Lo genera Odoo por numeración propia"),
        ("Retenciones del cliente", "Falta", "Contabilidad lo configura en la ficha del cliente"),
        ("Cuenta contable", "Falta", "Se define por categoría de producto en Odoo"),
        ("Cantidad y precio", "Lo da Smartier", "Campos `Cantidad` y `PrecioUnitario`"),
        ("Descuento", "Lo da Smartier", "Campo `Descuento`, porcentual"),
        ("Porcentaje de IVA", "Lo da Smartier", "Campo `PorcentajeIVA`: 16 % en los 50 productos"),
        ("Moneda", "Lo da Smartier", "`Nacional` → bolívares · `Extranjera` → dólares"),
        ("Fecha de la operación", "Lo da Smartier", "Campo `FechaEntregaReal`"),
        ("Tasa de cambio", "Lo pone Odoo", "Odoo aplica la tasa vigente al contabilizar"),
    ],
    anchos=(5.0, 3.2, 8.2),
    colores_col={1: color_estado},
)

subtitulo("El RIF: el bloqueo principal")

parrafo(
    "Es el único que impide facturar hoy. Odoo, con la normativa fiscal "
    "venezolana activa, **no permite emitir una factura sin identificación "
    "fiscal del cliente**. Lo comprobamos en las pruebas: Odoo rechazó la "
    "operación con un mensaje explícito hasta que completamos ese dato."
)

parrafo(
    "En Smartier el campo existe —se llama `Documento`— pero está **vacío en "
    "los tres clientes**. Consultamos también la ficha individual de cada uno, "
    "por si el detalle tuviera más información que el listado: no la tiene."
)

aviso("Cómo se puede resolver", [
    "**Opción preferible — cargarlo en Smartier.** El dato viaja solo y se "
    "mantiene actualizado. Conviene además hacerlo obligatorio al dar de alta "
    "un cliente nuevo, para que el problema no se repita.",
    "**Alternativa — cargarlo en Odoo.** Se introduce una vez por cliente y el "
    "middleware los enlaza. Válido si son pocos clientes, pero exige mantener "
    "dos sitios y decidir cómo se emparejan.",
    "Mientras no exista, el sistema **no descarta la nota en silencio**: la "
    "encola igualmente y la marca como pendiente con el motivo visible en el "
    "panel, para que alguien pueda corregirla.",
], tono="ambar")

subtitulo("Un punto a confirmar sobre el IVA")
parrafo(
    "Los 50 productos de Smartier tienen **16 %**. La instancia de Odoo con la "
    "que probamos tenía configurado **15 %**. Antes de la primera factura real "
    "hay que confirmar cuál es el vigente y dejarlo configurado en Odoo; si no, "
    "los totales no cuadrarán."
)


# ===========================================================================
# 06 · PRÓXIMOS PASOS
# ===========================================================================
seccion("06", "Próximos pasos")

PASOS = [
    ("Turicopy / Smartier", "Cargar el RIF de los tres clientes y decidir si "
     "pasa a ser obligatorio en el alta.", "Bloqueante"),
    ("Turicopy / Smartier", "Generar notas de entrega de prueba, en varios "
     "estados, para validar el circuito completo.", "Bloqueante"),
    ("Smart Automata", "Confirmar con Smartier el nombre exacto de los filtros, "
     "sobre todo el de facturación.", "Alta"),
    ("Contabilidad", "Confirmar si el IVA vigente es 16 % y definir la cuenta "
     "contable por categoría de producto.", "Alta"),
    ("Turicopy", "Disponer de una instancia de Odoo definitiva.", "Alta"),
    ("Contabilidad", "Indicar qué clientes son agentes de retención y en qué "
     "porcentaje.", "Media"),
    ("Turicopy", "Designar quién anota en Smartier el número de factura tras "
     "emitirla, para evitar duplicados.", "Media"),
]


def color_prioridad(valor):
    return ({"Bloqueante": GRANATE, "Alta": AMBAR}.get(valor, TENUE),
            valor in ("Bloqueante", "Alta"))


tabla_datos(
    ("Responsable", "Acción", "Prioridad"),
    [(r, a, p) for r, a, p in PASOS],
    anchos=(3.8, 9.4, 3.2),
    colores_col={2: color_prioridad},
)

parrafo(
    "En cuanto estén resueltos los dos primeros puntos, el circuito completo "
    "puede validarse **el mismo día**: el software ya está construido y "
    "probado en todas sus piezas.",
    espacio=4,
)


# ===========================================================================
# PIE
# ===========================================================================
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(18)
borde(p, "top", "E2DFD8", 6)
texto(p, "Informe preparado por Smart Automata sobre el middleware desarrollado "
         "para Turicopy. Las cifras y resultados proceden de pruebas ejecutadas "
         "contra la API de Smartier y contra Odoo, no de documentación. "
         "Verificado el 24 de agosto de 2026.", color=TENUE, tam=9)

SALIDA = r"c:\Users\SMARTAUT14\Desktop\API-Odoo\Informe-Integracion-Smartier-Odoo.docx"
doc.save(SALIDA)
print(f"Informe generado: {SALIDA}")
