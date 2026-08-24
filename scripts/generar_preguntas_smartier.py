"""
Genera el documento de preguntas para la reunion con Smartier en formato .docx.

Agenda tecnica: cada pregunta lleva su motivo y, donde la hay, la evidencia
medida contra la API real de Smartier (no lo que dice su documentacion).

Requiere python-docx, que NO esta en requirements.txt a proposito: la API no lo
necesita para funcionar, solo este script.
    pip install python-docx

Uso:  python scripts/generar_preguntas_smartier.py
El .docx resultante esta en .gitignore: es un entregable, se regenera con esto.
"""

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# Paleta: la misma del artifact (azul tecnico del propio PDF del cliente).
AZUL = RGBColor(0x1F, 0x4E, 0x79)
TINTA = RGBColor(0x16, 0x20, 0x2B)
TENUE = RGBColor(0x5A, 0x66, 0x72)
GRANATE = RGBColor(0x9B, 0x2C, 0x2C)
AMBAR = RGBColor(0x8A, 0x50, 0x0A)

doc = Document()

# --- Configuracion de pagina y estilo base -------------------------------
for seccion in doc.sections:
    seccion.top_margin = Cm(2.2)
    seccion.bottom_margin = Cm(2.2)
    seccion.left_margin = Cm(2.5)
    seccion.right_margin = Cm(2.5)

normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
normal.font.color.rgb = TINTA
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.15


def sombrear(parrafo, color_hex):
    """Aplica un fondo de color a un parrafo (Word no lo expone en la API)."""
    sombra = OxmlElement("w:shd")
    sombra.set(qn("w:val"), "clear")
    sombra.set(qn("w:fill"), color_hex)
    parrafo._p.get_or_add_pPr().append(sombra)


def borde_izquierdo(parrafo, color_hex, ancho=18):
    """Barra vertical de color a la izquierda del parrafo."""
    pbdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(ancho))
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), color_hex)
    pbdr.append(left)
    parrafo._p.get_or_add_pPr().append(pbdr)


def texto(parrafo, contenido, negrita=False, mono=False, color=None,
          cursiva=False, tam=None):
    """Anade un run con formato. Los tramos entre ` ` van en monoespaciada."""
    run = parrafo.add_run(contenido)
    run.bold = negrita
    run.italic = cursiva
    if mono:
        run.font.name = "Consolas"
        run.font.size = Pt(9.5)
    if tam:
        run.font.size = Pt(tam)
    if color:
        run.font.color.rgb = color
    return run


def parrafo_rico(p, contenido):
    """
    Escribe texto interpretando dos marcas: **negrita** y `monoespaciada`.

    Las marcas se pueden anidar ("**El `openapi.json`**"), asi que el interior
    de un tramo en negrita se vuelve a recorrer buscando codigo.
    """
    import re

    def escribir(fragmento, negrita=False):
        for sub in re.split(r"(`[^`]+`)", fragmento):
            if not sub:
                continue
            if sub.startswith("`") and sub.endswith("`"):
                texto(p, sub[1:-1], mono=True, color=AZUL, negrita=negrita)
            else:
                texto(p, sub, negrita=negrita)

    for trozo in re.split(r"(\*\*.+?\*\*)", contenido):
        if not trozo:
            continue
        if trozo.startswith("**") and trozo.endswith("**"):
            escribir(trozo[2:-2], negrita=True)
        else:
            escribir(trozo)


# --- Portada -------------------------------------------------------------
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(2)
texto(p, "TURICOPY · INTEGRACIÓN SMARTIER–ODOO", negrita=True, color=AZUL, tam=9)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(4)
run = texto(p, "Preguntas para Smartier", negrita=True, tam=26)
run.font.name = "Cambria"

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(10)
texto(p, "Agenda técnica de la reunión. Cada pregunta nace de algo que el "
         "middleware ya necesita resolver o de algo que se comprobó contra la "
         "API real y no cuadra con lo previsto.", color=TENUE, tam=11)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(14)
texto(p, "Base: ", negrita=True, tam=9)
texto(p, "turicopy.smartier.software/api/v2", mono=True, tam=9)
texto(p, "   ·   Verificado: ", negrita=True, tam=9)
texto(p, "21–24 ago 2026", tam=9)
texto(p, "   ·   Estado: ", negrita=True, tam=9)
texto(p, "adaptador construido, pendiente de datos", tam=9)

# --- Alerta destacada ----------------------------------------------------
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(6)
p.paragraph_format.space_after = Pt(2)
p.paragraph_format.left_indent = Cm(0.3)
sombrear(p, "FAEDED")
borde_izquierdo(p, "9B2C2C")
texto(p, "HALLAZGO QUE CONVIENE PLANTEAR PRIMERO", negrita=True,
      color=GRANATE, tam=9)

for frase in [
    "**La API ignora en silencio los filtros que no reconoce.** Enviamos "
    "`?FiltroQueNoExiste=x` y respondió `200` con normalidad, en lugar de "
    "rechazarlo. En cambio, un `Sort` inválido sí devuelve `400` con la lista "
    "de campos válidos.",
    "El riesgo es concreto: si escribimos mal `TieneFactura`, la sincronización "
    "traerá **todas** las notas —incluidas las ya facturadas— sin ningún aviso, "
    "y acabaríamos facturando dos veces. Todo el control de duplicados de la "
    "sección 8.2 del documento se apoya en ese filtro.",
]:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.space_after = Pt(4)
    sombrear(p, "FAEDED")
    borde_izquierdo(p, "9B2C2C")
    parrafo_rico(p, frase)


# --- Helpers de seccion --------------------------------------------------

def seccion(numero, titulo, nota=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(2)
    texto(p, f"{numero}  ", negrita=True, mono=True, color=AZUL, tam=11)
    run = texto(p, titulo, negrita=True, tam=15)
    run.font.name = "Cambria"
    # Linea inferior gruesa, como en el artifact.
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "3")
    bottom.set(qn("w:color"), "16202B")
    pbdr.append(bottom)
    p._p.get_or_add_pPr().append(pbdr)

    if nota:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        texto(p, nota, cursiva=True, color=TENUE, tam=10)


def pregunta(ref, titulo, porque, critica=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    texto(p, f"{ref}  ", negrita=True, mono=True,
          color=GRANATE if critica else TENUE, tam=9)
    texto(p, titulo, negrita=True, tam=11)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.75)
    parrafo_rico(p, porque)


def evidencia(lineas, aviso=False):
    fondo = "FBF1E0" if aviso else "EAF1F8"
    borde = "B4690E" if aviso else "1F4E79"
    color = AMBAR if aviso else AZUL

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.75)
    p.paragraph_format.space_after = Pt(1)
    sombrear(p, fondo)
    borde_izquierdo(p, borde, ancho=12)
    texto(p, "COMPROBADO" if not aviso else "COMPROBADO",
          negrita=True, color=color, tam=8)

    for linea, es_codigo in lineas:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.75)
        p.paragraph_format.space_after = Pt(3)
        sombrear(p, fondo)
        borde_izquierdo(p, borde, ancho=12)
        if es_codigo:
            texto(p, linea, mono=True, tam=9)
        else:
            parrafo_rico(p, linea)


# --- 01 Filtros ----------------------------------------------------------
seccion("01", "Filtros y control de duplicados",
        "Es el bloque más urgente: de aquí depende que no se facture dos veces.")

pregunta("1.1", "¿Cuál es el nombre exacto y el tipo del filtro TieneFactura?",
         "Necesitamos saber si espera `true`/`false`, `0`/`1`, o si el nombre "
         "lleva un prefijo. Y si hay forma de **validar** que el filtro se "
         "aplicó, más allá de confiar en que está bien escrito.", critica=True)
evidencia([("`?TieneFactura=false` y `?TieneFactura=true` devuelven ambos `200` "
            "con `Count: 0`. Con la base vacía no distinguimos si el filtro "
            "funciona o si se está ignorando.", False)], aviso=True)

pregunta("1.2", "¿Pueden validar los filtros desconocidos, como ya hacen con Sort?",
         "Un `400` ante un parámetro no reconocido convertiría un error "
         "silencioso —y caro— en uno visible al instante. El endpoint ya lo "
         "hace con el ordenamiento, así que el mecanismo existe.", critica=True)
evidencia([("Sort=CampoInexistente", True),
           ("→ 400 \"Campo de ordenamiento no válido.", True),
           ("   Campos ordenables: Id, Fecha, Estado.\"", True)])

pregunta("1.3", "¿Qué filtros de fecha acepta cada endpoint, con su nombre exacto?",
         "Sin filtro de fecha, cada pasada tendría que recorrer el histórico "
         "completo. Con el límite de 5 req/s, eso no escala.")
evidencia([("Probamos seis nombres —`FechaDesde`, `Desde`, `ModificadoDesde`, "
            "`FechaEntregaDesde` y otros— y **los seis** devuelven `200`. Como "
            "los desconocidos se ignoran, no podemos deducir cuál es el real.",
            False)], aviso=True)

pregunta("1.4", "¿Hay una marca de última modificación para sincronizar solo lo que cambió?",
         "Si una nota ya sincronizada cambia de precio o cantidad, ¿cómo nos "
         "enteramos? Los campos ordenables de notas son `Id`, `Fecha` y "
         "`Estado`: ninguno parece ser una fecha de modificación.")

# --- 02 RIF --------------------------------------------------------------
seccion("02", "El RIF de los clientes",
        "Bloqueo duro: sin identificación fiscal, Odoo no emite la factura.")

pregunta("2.1", "¿Por qué Documento.Contenido viene vacío en todos los clientes?",
         "¿Es que el campo no se usa en producción, que está vacío solo en este "
         "entorno, o que hace falta un permiso extra en la API Key para verlo?",
         critica=True)
evidencia([("GET /external/clientes/8", True),
           ("{\"Id\": 8, \"Nombre\": \"LISBETH SANCHEZ\",", True),
           (" \"Documento\": {\"Tipo\": 6, \"Contenido\": null}}", True),
           ("Los **3** clientes tienen `Contenido: null`. El detalle individual "
            "no aporta ningún campo que el listado no traiga ya.", False)],
          aviso=True)

pregunta("2.2", "¿Qué significa cada valor de Documento.Tipo?",
         "Los tres clientes tienen `Tipo: 6`. Necesitamos la tabla completa: "
         "qué es 6, y qué códigos corresponden a RIF (J, G, V) o a cédula. "
         "Odoo distingue el tipo de documento al emitir.")

pregunta("2.3", "¿Con qué formato llega el documento?",
         "¿`J-12345678-9`, `J123456789` o `12345678` con el prefijo por "
         "separado en `Tipo`? La deduplicación de la sección 6.2 se apoya en "
         "una coincidencia exacta: si el formato varía, falla.")

pregunta("2.4", "¿Puede Smartier exigir el documento como obligatorio al dar de alta un cliente?",
         "Es la solución de raíz. Si no, cada cliente nuevo sin RIF quedará "
         "bloqueado en el middleware hasta que alguien lo complete a mano.")

# --- 03 Notas ------------------------------------------------------------
seccion("03", "Notas de entrega",
        "Es la fuente de la factura, y es justo lo que no hemos podido ver nunca.")

pregunta("3.1", "¿Nos facilitan datos de prueba, o un ejemplo real de respuesta?",
         "Es lo más valioso que podemos sacar de la reunión. Con un solo JSON "
         "de ejemplo cerramos el mapeo entero; sin él, seguimos programando "
         "contra una estructura supuesta.", critica=True)
evidencia([("`notas-entrega`, `ordenes` y `tickets` devuelven `Count: 0`. Nunca "
            "hemos visto la forma real del DTO en una respuesta.", False)],
          aviso=True)

pregunta("3.2", "¿Una nota de entrega lleva un solo producto o varias líneas?",
         "El DTO muestra `Cantidad` y `PrecioUnitario` en **singular**, sin "
         "ningún array de líneas. Si es un producto por nota, hay que decidir "
         "si cada nota genera su propia factura o si se agrupan varias en una "
         "sola. Cambia el diseño por completo.", critica=True)

pregunta("3.3", "¿Qué contiene exactamente el campo Factura?",
         "¿Es el número que anotará contabilidad, un texto libre, un objeto? "
         "Nos serviría para conciliar contra el número que genere Odoo y "
         "detectar divergencias.")

pregunta("3.4", "¿Se factura Cantidad o CantidadTerminada?",
         "El DTO trae ambas. Si una entrega es parcial, ¿cuál manda? ¿Y puede "
         "una nota facturarse en dos tandas?")

pregunta("3.5", "¿Qué estados son facturables, y las de tipo Retiro también facturan?",
         "El enum tiene seis estados —`SinConfirmar`, `Pendiente`, "
         "`Coordinada`, `Facturada`, `EnTransito`, `Entregada`—. Hoy el "
         "middleware dispara con `Facturada`, pero conviene confirmarlo. Y el "
         "campo `Tipo` distingue `Retiro` de `Entrega`: ¿se tratan igual?")

# --- 04 Precios ----------------------------------------------------------
seccion("04", "Precios, moneda e impuestos")

pregunta("4.1", "¿Extranjera siempre significa dólar?",
         "El enum de moneda solo distingue `Nacional` y `Extranjera`. "
         "Traducimos a `VES` y `USD`, pero si algún día se factura en euros el "
         "modelo no lo distingue.")

pregunta("4.2", "¿El Monto del precio ya incluye IVA o es la base imponible?",
         "Determina si Odoo debe sumar el impuesto o desglosarlo. Un error "
         "aquí descuadra todas las facturas.")

pregunta("4.3", "¿El Descuento se aplica antes o después del IVA?",
         "Es un porcentaje sobre el precio unitario. Confirmamos el orden de "
         "cálculo para que el total cuadre al céntimo con la validación que ya "
         "hace el middleware.")

pregunta("4.4", "¿El 16 % es el único IVA vigente?",
         "Si aparece otra alícuota, necesitamos que exista el impuesto "
         "equivalente en Odoo antes de la primera factura con ese producto.")
evidencia([("Los **48** productos tienen `PorcentajeIVA: 16` y `Exento: false`. "
            "Ninguna otra alícuota, ningún exento.", False)])

pregunta("4.5", "¿Quién fija la tasa de cambio del día?",
         "Si Smartier guarda un precio en divisa y Odoo aplica su propia tasa, "
         "los importes divergen. Hace falta acordar cuál manda.")

# --- 05 Operacion --------------------------------------------------------
seccion("05", "Operación y soporte")

pregunta("5.1", "¿El límite de 5 req/s es por key o por tenant?",
         "Si es por key, podríamos pedir una segunda para separar la ingesta "
         "de las consultas puntuales. También: ¿devuelven cabecera "
         "`Retry-After` en el `429`? El cliente ya la respeta si viene.")

pregunta("5.2", "¿Hay entorno de pruebas separado del de producción?",
         "Para probar el flujo completo sin tocar datos reales, y para poder "
         "generar notas de prueba libremente.")

pregunta("5.3", "¿Nos pueden pasar el openapi.json?",
         "La documentación en pantalla exige sesión y el spec no está expuesto "
         "en las rutas habituales. Con el fichero trabajamos sin depender de "
         "su servidor y detectamos cambios de contrato.")

pregunta("5.4", "¿Cómo nos avisarán de cambios en la API?",
         "Un campo renombrado rompe la sincronización en silencio. ¿Hay "
         "versionado —el `v2` de la ruta— y aviso previo?")

pregunta("5.5", "Confirmación sobre las API Keys",
         "La clave que estamos usando conviene **revocarla y regenerarla** tras "
         "las pruebas. ¿Cuál es la duración recomendada y quién gestiona la "
         "rotación?")

# --- 06 Tabla de estado --------------------------------------------------
seccion("06", "Estado verificado de los datos",
        "Lo que devuelve la API hoy, medido el 21 y 24 de agosto de 2026.")

FILAS = [
    ("Clientes", "3", "Sin RIF", "Id, Nombre, CreadoUtc, Tipo, Estado"),
    ("Productos", "48", "Completos", "Id, Nombre, Estado, Tipo"),
    ("Notas de entrega", "0", "Vacío", "Id, Fecha, Estado"),
    ("Órdenes", "0", "Vacío", "Id, FechaEntrega, Prioridad"),
    ("Vendedores", "3", "Opcional", "—"),
    ("Sectores", "24", "No contable", "—"),
    ("Recursos", "82", "No contable", "—"),
]

tabla = doc.add_table(rows=1, cols=4)
tabla.style = "Table Grid"
tabla.alignment = WD_TABLE_ALIGNMENT.CENTER

encabezados = ("Recurso", "Registros", "Situación", "Campos ordenables")
for i, cab in enumerate(encabezados):
    celda = tabla.rows[0].cells[i]
    celda.text = ""
    p = celda.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    texto(p, cab.upper(), negrita=True, color=AZUL, tam=8)
    sombrear(p, "EAF1F8")

COLOR_ESTADO = {"Sin RIF": GRANATE, "Vacío": GRANATE, "Completos": AZUL}

for recurso, num, situacion, orden in FILAS:
    celdas = tabla.add_row().cells
    for i, (valor, mono) in enumerate([
        (recurso, False), (num, True), (situacion, False), (orden, True)
    ]):
        celdas[i].text = ""
        p = celdas[i].paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        if i == 2:
            texto(p, valor, negrita=True,
                  color=COLOR_ESTADO.get(valor, TENUE), tam=9)
        else:
            texto(p, valor, mono=mono, tam=9)

# --- Cierre --------------------------------------------------------------
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(20)
p.paragraph_format.space_after = Pt(6)
run = texto(p, "Qué pedir antes de salir de la reunión", negrita=True, tam=14)
run.font.name = "Cambria"
pbdr = OxmlElement("w:pBdr")
top = OxmlElement("w:top")
top.set(qn("w:val"), "single")
top.set(qn("w:sz"), "12")
top.set(qn("w:space"), "6")
top.set(qn("w:color"), "16202B")
pbdr.append(top)
p._p.get_or_add_pPr().append(pbdr)

PEDIDOS = [
    "**Un JSON real de nota de entrega.** Con eso se cierra el mapeo completo. "
    "Es lo que más desbloquea.",
    "**El `openapi.json`** o la lista de filtros válidos por endpoint, con "
    "nombre y tipo exactos.",
    "**Un compromiso sobre el RIF:** quién lo carga, cuándo, y si pasa a ser "
    "obligatorio en el alta.",
    "**Datos de prueba** con notas en varios estados, para validar el flujo de "
    "punta a punta.",
    "**La tabla de `Documento.Tipo`** y el formato con que llega el número.",
]

for i, pedido in enumerate(PEDIDOS, start=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.left_indent = Cm(0.9)
    p.paragraph_format.first_line_indent = Cm(-0.9)
    texto(p, f"{i:02d}   ", negrita=True, mono=True, color=GRANATE, tam=9)
    parrafo_rico(p, pedido)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(16)
pbdr = OxmlElement("w:pBdr")
top = OxmlElement("w:top")
top.set(qn("w:val"), "single")
top.set(qn("w:sz"), "6")
top.set(qn("w:space"), "6")
top.set(qn("w:color"), "E2DFD8")
pbdr.append(top)
p._p.get_or_add_pPr().append(pbdr)
texto(p, "Preparado por Smart Automata a partir del documento de integración y "
         "de pruebas ejecutadas contra la API de Smartier. Las respuestas "
         "marcadas como ", color=TENUE, tam=9)
texto(p, "Comprobado", cursiva=True, color=TENUE, tam=9)
texto(p, " proceden de llamadas reales, no de la documentación.",
      color=TENUE, tam=9)

SALIDA = r"c:\Users\SMARTAUT14\Desktop\API-Odoo\Preguntas-Smartier-Reunion.docx"
doc.save(SALIDA)
print(f"Documento generado: {SALIDA}")
