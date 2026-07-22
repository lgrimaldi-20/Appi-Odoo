from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

# ─── Estilos globales ───
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)

# Márgenes
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)


def add_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0, 51, 102)


def add_subtitle(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0, 76, 153)
    p.space_before = Pt(12)


def add_section_header(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0, 102, 51)
    p.space_before = Pt(8)


def add_text(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)
    p.space_after = Pt(2)
    return p


def add_equation(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'Cambria Math'
    p.space_before = Pt(4)
    p.space_after = Pt(4)
    return p


def add_result(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("  ➤  " + text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(153, 0, 0)
    p.space_before = Pt(4)
    p.space_after = Pt(6)


def add_dato(text):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(10)
    p.space_after = Pt(1)


def add_line():
    p = doc.add_paragraph()
    run = p.add_run("─" * 70)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(150, 150, 150)
    p.space_before = Pt(6)
    p.space_after = Pt(6)


# ═══════════════════════════════════════════════════════════════
# PORTADA
# ═══════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
add_title("TRANSFERENCIA DE CALOR")
add_title("Resolución de Ejercicios")
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Aletas y Conducción Transitoria")
run.font.size = Pt(13)
run.italic = True
doc.add_paragraph()
doc.add_paragraph()

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# EJERCICIO 1
# ═══════════════════════════════════════════════════════════════
add_title("EJERCICIO 1")
add_subtitle("Aleta rectangular de cobre — Determinación de temperatura base, calor y eficiencia")

add_section_header("Datos del problema:")
add_dato("k = 387.6 W/(m·°C)  —  Conductividad térmica del cobre")
add_dato("P = 15 cm = 0.15 m  —  Perímetro de la sección transversal")
add_dato("A_c = 10 cm² = 1 × 10⁻³ m²  —  Área de sección transversal")
add_dato("L = 12 cm = 0.12 m  —  Longitud de la aleta")
add_dato("h = 25 W/(m²·K)  —  Coeficiente convectivo")
add_dato("T∞ = 32°C  —  Temperatura del ambiente")
add_dato("T(x₁ = 3 cm) = 250°C  —  Temperatura a 3 cm de la base")
add_dato("T(x₂ = 7 cm) = 200°C  —  Temperatura a 7 cm de la base")

add_section_header("Cálculo del parámetro m:")
add_equation("m = √(h·P / k·A_c)")
add_equation("m = √(25 × 0.15 / 387.6 × 1×10⁻³)")
add_equation("m = √(3.75 / 0.3876) = √9.675")
add_equation("m = 3.1105 m⁻¹")

add_line()
add_section_header("a) Temperatura de la base T_b")

add_text("Usando la solución general de la ecuación de aleta:")
add_equation("θ(x) = C₁·e^(mx) + C₂·e^(-mx)")
add_text("Donde θ(x) = T(x) - T∞")
add_text("")
add_text("Excesos de temperatura:")
add_dato("θ(0.03 m) = 250 - 32 = 218 °C")
add_dato("θ(0.07 m) = 200 - 32 = 168 °C")

add_text("")
add_text("Planteando el sistema de ecuaciones:")
add_equation("218 = C₁·e^(3.1105×0.03) + C₂·e^(-3.1105×0.03)")
add_equation("218 = 1.09778·C₁ + 0.91094·C₂   ... (Ec. 1)")
add_equation("")
add_equation("168 = C₁·e^(3.1105×0.07) + C₂·e^(-3.1105×0.07)")
add_equation("168 = 1.24332·C₁ + 0.80430·C₂   ... (Ec. 2)")

add_text("")
add_text("Resolviendo el sistema:")
add_equation("De Ec.1:  C₁ = (218 - 0.91094·C₂) / 1.09778")
add_text("Sustituyendo en Ec.2:")
add_equation("168 = [(218 - 0.91094·C₂)/1.09778] × 1.24332 + 0.80430·C₂")
add_equation("168 = 246.834 - 1.03235·C₂ + 0.80430·C₂")
add_equation("168 = 246.834 - 0.22805·C₂")
add_equation("0.22805·C₂ = 78.834")
add_equation("C₂ = 345.69 °C")
add_equation("")
add_equation("C₁ = (218 - 0.91094 × 345.69) / 1.09778 = -88.27 °C")

add_text("")
add_text("Temperatura de la base (x = 0):")
add_equation("θ_b = C₁ + C₂ = -88.27 + 345.69 = 257.42 °C")
add_equation("T_b = θ_b + T∞ = 257.42 + 32")
add_result("T_b = 289.42 °C")

add_line()
add_section_header("b) Pérdida de calor de la aleta (W)")

add_equation("q_aleta = √(h·P·k·A_c) · θ_b · tanh(mL)")
add_equation("√(25 × 0.15 × 387.6 × 10⁻³) = √1.4535 = 1.2056 W/°C")
add_equation("tanh(3.1105 × 0.12) = tanh(0.37326) = 0.3575")
add_equation("q_aleta = 1.2056 × 257.42 × 0.3575")
add_result("q_aleta = 110.89 W")

add_line()
add_section_header("c) Eficiencia de la aleta (%)")

add_equation("η = tanh(mL) / (mL)")
add_equation("η = 0.3575 / 0.37326")
add_result("η = 95.76 %")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# EJERCICIO 2
# ═══════════════════════════════════════════════════════════════
add_title("EJERCICIO 2")
add_subtitle("Arreglo de aletas rectangulares de cobre en pared")

add_section_header("Datos del problema:")
add_dato("t = 1 mm = 0.001 m  —  Espesor de la aleta")
add_dato("L = 10 mm = 0.01 m  —  Longitud de la aleta")
add_dato("k = 380 W/(m·°C)  —  Conductividad térmica del cobre")
add_dato("T_b = 230°C  —  Temperatura de la pared (base)")
add_dato("H = 1 m  —  Altura de la pared")
add_dato("W = 1 m  —  Ancho de la pared")
add_dato("T∞ = 30°C  —  Temperatura ambiente")
add_dato("h = 40 W/(m²·°C)  —  Coeficiente de transferencia de calor")
add_dato("S = 8 mm = 0.008 m  —  Espaciado entre aletas")
add_dato("θ_b = T_b - T∞ = 230 - 30 = 200 °C  —  Exceso de temperatura")

add_section_header("Parámetros geométricos:")
add_text("Para aleta rectangular delgada (ancho W = 1 m):")
add_equation("P = 2(W + t) = 2(1 + 0.001) ≈ 2.002 m")
add_equation("A_c = W × t = 1 × 0.001 = 0.001 m²")
add_equation("L_c = L + t/2 = 0.01 + 0.0005 = 0.0105 m  (longitud corregida)")

add_section_header("Parámetro m:")
add_equation("m = √(h·P / k·A_c) = √(40 × 2.002 / 380 × 0.001)")
add_equation("m = √(80.08 / 0.38) = √210.74 = 14.517 m⁻¹")
add_equation("mL_c = 14.517 × 0.0105 = 0.15243")

add_line()
add_section_header("a) Eficiencia de la aleta (%)")

add_equation("η_aleta = tanh(mL_c) / (mL_c)")
add_equation("tanh(0.15243) = 0.15126")
add_equation("η_aleta = 0.15126 / 0.15243")
add_result("η_aleta = 99.23 %")

add_line()
add_section_header("b) Eficiencia global de las aletas (%)")

add_text("Número de aletas:")
add_equation("Paso = S + t = 8 + 1 = 9 mm = 0.009 m")
add_equation("n = H / Paso = 1 / 0.009 ≈ 111 aletas")

add_text("")
add_text("Cálculo de áreas:")
add_dato("Área de una aleta: A_aleta = 2 × W × L_c = 2 × 1 × 0.0105 = 0.021 m²")
add_dato("Área total de aletas: n × A_aleta = 111 × 0.021 = 2.331 m²")
add_dato("Área total de la pared: A_total = H × W = 1 × 1 = 1 m²")
add_dato("Área sin aletas: A_sin_aletas = 1 - 111 × 0.001 × 1 = 0.889 m²")
add_dato("Área total expuesta: A_T = 0.889 + 2.331 = 3.220 m²")

add_text("")
add_equation("η_global = 1 - (n·A_aleta / A_T) × (1 - η_aleta)")
add_equation("η_global = 1 - (2.331/3.220) × (1 - 0.9923)")
add_equation("η_global = 1 - 0.7239 × 0.0077 = 1 - 0.00557")
add_result("η_global = 99.44 %")

add_line()
add_section_header("c) Transferencia de calor con arreglo de aletas (W)")

add_equation("q_total = η_global × h × A_T × θ_b")
add_equation("q_total = 0.9944 × 40 W/(m²·°C) × 3.220 m² × 200 °C")
add_result("q_total = 25,598.2 W ≈ 25.6 kW")

add_line()
add_section_header("d) Transferencia de calor de la pared sin aletas (W)")

add_equation("q_sin_aletas = h × A_pared × θ_b")
add_equation("q_sin_aletas = 40 W/(m²·°C) × 1 m² × 200 °C")
add_result("q_sin_aletas = 8,000 W = 8 kW")

add_text("")
add_text("Las aletas incrementan la transferencia de calor en un factor de 25,598 / 8,000 = 3.2 veces.", bold=True)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# EJERCICIO 3
# ═══════════════════════════════════════════════════════════════
add_title("EJERCICIO 3")
add_subtitle("Cilindro largo — Conducción transitoria (Método aproximado)")

add_section_header("Datos del problema:")
add_dato("r₀ = 20 cm = 0.20 m  —  Radio del cilindro")
add_dato("k = 170 W/(m·K)  —  Conductividad térmica")
add_dato("α = 9.05 × 10⁻⁷ m²/s  —  Difusividad térmica")
add_dato("T_i = 650°C  —  Temperatura inicial uniforme")
add_dato("T∞ = 75°C  —  Temperatura del medio (sustancia)")
add_dato("h = 1700 W/(m²·K)  —  Coeficiente de convección")
add_dato("T(r,t) = 250°C a profundidad de 20 mm → r = r₀ - 0.02 = 0.18 m")

add_section_header("Número de Biot:")
add_equation("Bi = h·r₀ / k = (1700 × 0.20) / 170 = 340/170")
add_equation("Bi = 2.0")
add_text("Como Bi > 0.1, NO se puede usar capacitancia global → Se usa método aproximado (un término).", bold=True)

add_section_header("Parámetros de tablas para Bi = 2.0 (cilindro):")
add_dato("ζ₁ = 1.5995 rad")
add_dato("C₁ = 1.3384")

add_text("")
add_text("Solución de un término para cilindro largo infinito:")
add_equation("θ* = (T(r,t) - T∞) / (T_i - T∞) = C₁ · e^(-ζ₁²·Fo) · J₀(ζ₁ · r/r₀)")

add_line()
add_section_header("Encontrar el tiempo t:")

add_text("En el punto r = 0.18 m → r/r₀ = 0.18/0.20 = 0.9")
add_equation("θ* = (250 - 75) / (650 - 75) = 175/575 = 0.30435")

add_text("")
add_text("Función de Bessel:")
add_equation("J₀(ζ₁ × r/r₀) = J₀(1.5995 × 0.9) = J₀(1.43955) ≈ 0.5531")

add_text("")
add_text("Sustituyendo:")
add_equation("0.30435 = 1.3384 × e^(-(1.5995)²·Fo) × 0.5531")
add_equation("0.30435 = 0.7403 × e^(-2.5584·Fo)")
add_equation("e^(-2.5584·Fo) = 0.30435 / 0.7403 = 0.41111")
add_equation("-2.5584·Fo = ln(0.41111) = -0.88916")
add_equation("Fo = 0.88916 / 2.5584 = 0.34754")

add_text("")
add_text("Cálculo del tiempo:")
add_equation("Fo = α·t / r₀²  →  t = Fo·r₀² / α")
add_equation("t = (0.34754 × (0.20)²) / (9.05 × 10⁻⁷)")
add_equation("t = 0.013902 / 9.05 × 10⁻⁷")
add_result("t = 15,361 s ≈ 4 horas 16 minutos")

add_line()
add_section_header("Temperatura en el eje central (r = 0):")

add_text("En el centro, J₀(0) = 1:")
add_equation("θ₀* = C₁ · e^(-ζ₁²·Fo) = 1.3384 × 0.41111 = 0.55015")
add_equation("T(0,t) = T∞ + θ₀* × (T_i - T∞)")
add_equation("T(0,t) = 75 + 0.55015 × 575")
add_result("T_centro = 391.3 °C")

add_line()
add_section_header("Calor transferido por unidad de longitud:")

add_text("Calor máximo por unidad de longitud:")
add_equation("ρ·c_p = k / α = 170 / (9.05 × 10⁻⁷) = 1.8785 × 10⁸ J/(m³·K)")
add_equation("Q_max/L = ρ·c_p · π · r₀² · (T_i - T∞)")
add_equation("Q_max/L = 1.8785×10⁸ × π × (0.20)² × 575")
add_equation("Q_max/L = 1.3573 × 10¹⁰ J/m")

add_text("")
add_text("Razón de calor transferido:")
add_equation("J₁(1.5995) ≈ 0.5699  (de tablas)")
add_equation("Q/Q_max = 1 - 2·θ₀* · J₁(ζ₁)/ζ₁")
add_equation("Q/Q_max = 1 - 2 × 0.55015 × (0.5699/1.5995)")
add_equation("Q/Q_max = 1 - 0.39203 = 0.60797")

add_text("")
add_equation("Q/L = 0.60797 × 1.3573 × 10¹⁰")
add_result("Q/L = 8,252 MJ/m")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# TABLA RESUMEN
# ═══════════════════════════════════════════════════════════════
add_title("TABLA RESUMEN DE RESULTADOS")
doc.add_paragraph()

table = doc.add_table(rows=11, cols=3)
table.style = 'Medium Shading 1 Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ["Ejercicio", "Inciso", "Resultado"]
for i, h_text in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h_text
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.bold = True

data = [
    ("1", "a) T_base", "289.42 °C"),
    ("1", "b) q_aleta", "110.89 W"),
    ("1", "c) η_aleta", "95.76 %"),
    ("2", "a) η_aleta", "99.23 %"),
    ("2", "b) η_global", "99.44 %"),
    ("2", "c) q_con aletas", "25,598.2 W"),
    ("2", "d) q_sin aletas", "8,000 W"),
    ("3", "Tiempo t", "15,361 s ≈ 4.27 h"),
    ("3", "T_centro", "391.3 °C"),
    ("3", "Q/L", "8,252 MJ/m"),
]

for row_idx, (ej, inc, res) in enumerate(data, start=1):
    table.rows[row_idx].cells[0].text = ej
    table.rows[row_idx].cells[1].text = inc
    table.rows[row_idx].cells[2].text = res
    for cell in table.rows[row_idx].cells:
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Guardar
output_path = os.path.join(os.path.dirname(__file__), "..", "Transferencia_de_Calor_Ejercicios.docx")
output_path = os.path.normpath(output_path)
doc.save(output_path)
print(f"Documento guardado en: {output_path}")
